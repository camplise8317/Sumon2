import streamlit as st
import pandas as pd
import google.generativeai as genai
import PyPDF2
import docx
import re
import io # Importar el módulo io para manejar archivos en memoria
import openai # Importar la librería OpenAI para modelos GPT

# --- Configuración de la API de Gemini y OpenAI ---
# Se recomienda usar st.secrets para API keys en despliegues reales
# Para pruebas en Colab, se puede usar st.sidebar.text_input.

st.sidebar.header("Configuración de API Keys")
gemini_api_key = st.sidebar.text_input("API Key de Google Gemini", type="password", help="Obtén tu clave en https://aistudio.google.com/app/apikey")
openai_api_key = st.sidebar.text_input("API Key de OpenAI (para modelos GPT)", type="password", help="Obtén tu clave en https://platform.openai.com/account/api-keys")

# Inicialización condicional de Gemini y OpenAI
gemini_config_ok = False
openai_config_ok = False

if gemini_api_key:
    try:
        genai.configure(api_key=gemini_api_key)
        gemini_config_ok = True
        st.sidebar.success("API Key de Gemini configurada.")
    except Exception as e:
        st.sidebar.error(f"Error al configurar la API Key de Gemini: {e}")
else:
    st.sidebar.warning("Por favor, ingresa tu API Key de Gemini para usar modelos Gemini.")

if openai_api_key:
    openai.api_key = openai_api_key
    openai_config_ok = True
    st.sidebar.success("API Key de OpenAI configurada.")
else:
    st.sidebar.warning("Por favor, ingresa tu API Key de OpenAI para usar modelos GPT.")

# --- Funciones de Lectura de Archivos (Adaptadas para Streamlit Uploader) ---
@st.cache_data # Decorador de Streamlit para cachear los datos y no recargar el Excel/PDF cada vez
def leer_excel_cargado(uploaded_file):
    """
    Lee un archivo Excel cargado por Streamlit y lo carga en un DataFrame de pandas.
    """
    if uploaded_file is not None:
        try:
            df = pd.read_excel(uploaded_file)
            st.sidebar.success(f"Archivo Excel '{uploaded_file.name}' cargado exitosamente.")
            return df
        except Exception as e:
            st.sidebar.error(f"Ocurrió un error al leer el archivo Excel: {e}")
            return None
    return None

@st.cache_data # Decorador de Streamlit para cachear el texto del PDF
def leer_pdf_cargado(uploaded_file):
    """
    Lee el texto de un archivo PDF cargado por Streamlit.
    """
    if uploaded_file is not None:
        try:
            texto_pdf = ""
            # PyPDF2.PdfReader necesita un objeto tipo archivo binario
            reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
            for page_num in range(len(reader.pages)):
                texto_pdf += reader.pages[page_num].extract_text()
            st.sidebar.success(f"Archivo PDF '{uploaded_file.name}' leído exitosamente.")
            return texto_pdf
        except Exception as e:
            st.sidebar.error(f"Ocurrió un error al leer el archivo PDF: {e}")
            return ""
    return ""

# --- Función para obtener la descripción de la taxonomía de Bloom ---
def get_descripcion_bloom(proceso_cognitivo_elegido):
    descripcion_bloom_map = {
        "RECORDAR": "Recuperar información relevante desde la memoria de largo plazo.",
        "COMPRENDER": "Construir significado a partir de información mediante interpretación, resumen, explicación u otras tareas.",
        "APLICAR": "Usar procedimientos en situaciones conocidas o nuevas.",
        "ANALIZAR": "Descomponer información y examinar relaciones entre partes.",
        "EVALUAR": "Emitir juicios basados en criterios para valorar ideas o soluciones.",
        "CREAR": "Generar nuevas ideas, productos o formas de reorganizar información."
    }
    return descripcion_bloom_map.get(str(proceso_cognitivo_elegido).upper(), "Descripción no disponible para este proceso cognitivo.")

# --- Función para generar texto con Gemini o GPT ---
def generar_texto_con_llm(model_type, model_name, prompt):
    if model_type == "Gemini":
        if not gemini_config_ok:
            st.error("API Key de Gemini no configurada. No se puede generar texto con Gemini.")
            return None
        modelo = genai.GenerativeModel(model_name)
        response = modelo.generate_content(prompt)
        return response.text
    elif model_type == "GPT":
        if not openai_config_ok:
            st.error("API Key de OpenAI no configurada. No se puede generar texto con GPT.")
            return None
        client = openai.OpenAI(api_key=openai.api_key)
        response = client.chat.completions.create(
            model=model_name,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=2000 # Ajusta según necesidad
        )
        return response.choices[0].message.content
    return None

# --- Función para auditar el ítem generado ---
def auditar_item_con_llm(model_type, model_name, item_generado, grado, area, asignatura, estacion, 
                         proceso_cognitivo, nanohabilidad, microhabilidad, 
                         competencia_nanohabilidad, contexto_educativo, manual_reglas_texto="", descripcion_bloom="", grafico_necesario="", descripcion_grafico=""):
    """
    Audita un ítem generado para verificar su cumplimiento con criterios específicos.
    """
    auditoria_prompt = f"""
    Eres un experto en validación de ítems educativos, especializado en pruebas tipo ICFES y las directrices del equipo IMPROVE.
    Tu tarea es AUDITAR RIGUROSAMENTE el siguiente ítem generado por un modelo de lenguaje.

    Debes verificar que el ítem cumpla con TODOS los siguientes criterios, prestando especial atención a la alineación con los parámetros proporcionados y a las reglas de formato y contenido.

    --- CRITERIOS DE AUDITORÍA ---
    1.  **Formato del Enunciado:** ¿El enunciado está formulado como pregunta clara y directa, sin ambigüedades ni errores?
    2.  **Número de Opciones:** ¿Hay exactamente 3 opciones (A, B, C)?
    3.  **Respuesta Correcta Indicada:** ¿La sección 'RESPUESTA CORRECTA:' está claramente indicada y coincide con una de las opciones?
    4.  **Diseño de Justificaciones:**
        * ¿Hay justificaciones bien diferenciadas para CADA opción (A, B, C)?
        * ¿La justificación de la opción **correcta** explica el razonamiento, procedimiento o estrategia relevante (NO por descarte)?
        * ¿Las justificaciones de las opciones **incorrectas** están redactadas siguiendo el formato: “El estudiante podría escoger la opción X porque… Sin embargo, esto es incorrecto porque…”?
    5.  **Estilo y Restricciones:** ¿No se usan negaciones mal redactadas, nombres reales, marcas, lugares reales, datos personales o frases vagas como “ninguna de las anteriores” o “todas las anteriores”?
    6.  **Alineación del Contenido:** ¿El ítem (contexto, enunciado, opciones) está alineado EXCLUSIVAMENTE con los siguientes elementos temáticos y cognitivos?
        * Grado: {grado}
        * Área: {area}
        * Asignatura: {asignatura}
        * Estación o unidad temática: {estacion}
        * Proceso Cognitivo (Taxonomía de Bloom): {proceso_cognitivo} (su descripción es "{descripcion_bloom}")
        * Nanohabilidad (foco principal): {nanohabilidad}
        * Microhabilidad (evidencia de aprendizaje): {microhabilidad}
        * Competencia (asociada a Nanohabilidad): {competencia_nanohabilidad}
        * Nivel educativo del estudiante: {contexto_educativo}
    7.  **Gráfico (si aplica):** Si el ítem indica que requiere un gráfico, ¿la descripción del gráfico es clara, detallada y funcional para su futura creación?
        * Gráfico Necesario: {grafico_necesario}
        * Descripción del Gráfico: {descripcion_grafico if grafico_necesario == 'SÍ' else 'N/A'}

    --- MANUAL DE REGLAS ADICIONAL ---
    Las siguientes reglas son de suma importancia para la calidad y pertinencia del ítem. Debes asegurar que el ítem cumple con todas ellas.
    {manual_reglas_texto}
    -----------------------------------

    ÍTEM A AUDITAR:
    --------------------
    {item_generado}
    --------------------

    Devuelve tu auditoría con este formato estructurado:

    VALIDACIÓN DE CRITERIOS:
    - Formato del Enunciado: [✅ / ❌] + Comentario (si ❌)
    - Número de Opciones (3): [✅ / ❌]
    - Respuesta Correcta Indicada: [✅ / ❌]
    - Diseño de Justificaciones: [✅ / ⚠️ / ❌] + Observaciones (si ⚠️/❌)
    - Estilo y Restricciones: [✅ / ⚠️ / ❌] + Observaciones (si ⚠️/❌)
    - Alineación del Contenido: [✅ / ❌] + Comentario (si ❌)
    - Gráfico (si aplica): [✅ / ⚠️ / ❌] + Observaciones (si ⚠️/❌)

    DICTAMEN FINAL:
    [✅ CUMPLE TOTALMENTE / ⚠️ CUMPLE PARCIALMENTE / ❌ RECHAZADO]

    OBSERVACIONES FINALES:
    [Explica de forma concisa qué aspectos necesitan mejora, si el dictamen no es ✅. Si es ✅, puedes indicar "El ítem cumple con todos los criterios."]
    """
    return generar_texto_con_llm(model_type, model_name, auditoria_prompt)

# --- Función para generar preguntas usando el modelo de generación seleccionado ---
def generar_pregunta_con_seleccion(gen_model_type, gen_model_name, audit_model_type, audit_model_name, 
                                 fila_datos, criterios_generacion, manual_reglas_texto="", informacion_adicional_usuario=""):
    """
    Genera una pregunta educativa de opción múltiple usando el modelo de generación seleccionado
    y la itera para refinarla si la auditoría lo requiere.
    """
    tipo_pregunta = criterios_generacion.get("tipo_pregunta", "opción múltiple con 3 opciones") 
    dificultad = criterios_generacion.get("dificultad", "media")
    contexto_educativo = criterios_generacion.get("contexto_educativo", "general")
    formato_justificacion = criterios_generacion.get("formato_justificacion", """
        • Justificación correcta: debe explicar el razonamiento o proceso cognitivo (NO por descarte).
        • Justificaciones incorrectas: deben redactarse como: “El estudiante podría escoger la opción X porque… Sin embargo, esto es incorrecto porque…”
    """)
    
    grado_elegido = fila_datos.get('GRADO', 'no especificado')
    area_elegida = fila_datos.get('ÁREA', 'no especificada')
    asignatura_elegida = fila_datos.get('ASIGNATURA', 'no especificada')
    estacion_elegida = fila_datos.get('ESTACIÓN', 'no especificada')
    proceso_cognitivo_elegido = fila_datos.get('PROCESO COGNITIVO', 'no especificado')
    nanohabilidad_elegida = fila_datos.get('NANOHABILIDAD', 'no especificada')
    microhabilidad_elegida = fila_datos.get('MICROHABILIDAD', 'no especificada')
    competencia_nanohabilidad_elegida = fila_datos.get('COMPETENCIA NANOHABILIDAD', 'no especificada')

    dato_para_pregunta_foco = nanohabilidad_elegida
    descripcion_bloom = get_descripcion_bloom(proceso_cognitivo_elegido)

    current_item_text = ""
    auditoria_status = "❌ RECHAZADO" # Estado inicial
    audit_observations = "" # Observaciones para el refinamiento
    max_refinement_attempts = 5 # Número máximo de intentos de refinamiento
    attempt = 0
    grafico_necesario = "NO" # Valor por defecto
    descripcion_grafico = "" # Valor por defecto

    # Almacenar detalles de clasificación para el ítem
    classification_details = {
        "Grado": grado_elegido,
        "Área": area_elegida,
        "Asignatura": asignatura_elegida,
        "Estación": estacion_elegida,
        "Proceso Cognitivo": proceso_cognitivo_elegido,
        "Nanohabilidad": nanohabilidad_elegida,
        "Microhabilidad": microhabilidad_elegida,
        "Competencia Nanohabilidad": competencia_nanohabilidad_elegida
    }

    item_final_data = None # Variable para guardar el ítem final (aprobado o la última versión auditada)

    while auditoria_status != "✅ CUMPLE TOTALMENTE" and attempt < max_refinement_attempts:
        attempt += 1
        st.info(f"--- Generando/Refinando Ítem (Intento {attempt}/{max_refinement_attempts}) ---")

        prompt_content_for_llm = f"""
        Eres un diseñador experto en ítems de evaluación educativa, especializado en pruebas tipo ICFES u otras de alta calidad técnica.

        Tu tarea es construir un ítem de {tipo_pregunta} con una única respuesta correcta, cumpliendo rigurosamente las reglas de construcción de ítems y alineado con el marco cognitivo de la Taxonomía de Bloom.

        --- CONTEXTO Y PARÁMETROS DEL ÍTEM ---
        - Grado: {grado_elegido}
        - Área: {area_elegida}
        - Asignatura: {asignatura_elegida}
        - Estación o unidad temática: {estacion_elegida}
        - Proceso cognitivo (Taxonomía de Bloom): {proceso_cognitivo_elegido}
        - Descripción del proceso cognitivo:
          "{descripcion_bloom}"
        - Nanohabilidad (foco principal del ítem): {nanohabilidad_elegida}
        - Nivel educativo esperado del estudiante: {contexto_educativo}
        - Nivel de dificultad deseado: {dificultad}

        --- INSTRUCCIONES PARA LA CONSTRUCCIÓN DEL ÍTEM ---
        CONTEXTO DEL ÍTEM:
        - Incluye una situación contextualizada, relevante y plausible para el grado y área indicada.
        - La temática debe ser la de la {estacion_elegida}, y esto debe ser central, no una mera contextualización.
        - La situación debe ser funcional: debe activar el pensamiento requerido por la nanohabilidad.
        - Debe garantizarse que el proceso cognitivo corresponde fielmente a la descripción de la taxonomia de Bloom.
        - Evita referencias a marcas, nombres propios, lugares reales o información personal identificable.

        ENUNCIADO:
        - Formula una pregunta clara, directa, sin ambigüedades ni tecnicismos innecesarios.
        - Si utilizas negaciones, resáltalas en MAYÚSCULAS Y NEGRITA (por ejemplo: **NO ES**, **EXCEPTO**).
        - Asegúrate de que el enunciado refleje el tipo de tarea cognitiva esperado según el proceso de Bloom.

        OPCIONES DE RESPUESTA:
        - Escribe exactamente tres opciones (A, B y C).
        - Solo una opción debe ser correcta.
        - Los distractores (respuestas incorrectas) deben estar bien diseñados: deben ser creíbles, funcionales y representar errores comunes o concepciones alternativas frecuentes.
        - No utilices fórmulas vagas como “ninguna de las anteriores” o “todas las anteriores”.

        JUSTIFICACIONES:
        {formato_justificacion}

        --- REGLAS ADICIONALES DEL MANUAL DE CONSTRUCCIÓN ---
        Considera y aplica estrictamente todas las directrices, ejemplos y restricciones contenidas en el siguiente manual.
        Esto es de suma importancia para la calidad y pertinencia del ítem.

        Manual de Reglas:
        {manual_reglas_texto}
        ----------------------------------------------------

        --- INFORMACIÓN ADICIONAL PROPORCIONADA POR EL USUARIO ---
        {informacion_adicional_usuario if informacion_adicional_usuario else "No se proporcionó información adicional."}
        ----------------------------------------------------------

        --- DATO CLAVE PARA LA CONSTRUCCIÓN ---
        Basado en el foco temático y el proceso cognitivo, considera el siguiente dato o idea esencial:
        "{dato_para_pregunta_foco}"

        --- INSTRUCCIONES ESPECÍFICAS DE SALIDA PARA GRÁFICO ---
        Después del bloque de JUSTIFICACIONES, incluye la siguiente información para indicar si el ítem necesita un gráfico y cómo sería:
        GRAFICO_NECESARIO: [SÍ/NO]
        DESCRIPCION_GRAFICO: [Si GRAFICO_NECESARIO es SÍ, proporciona una descripción MUY DETALLADA del gráfico. Incluye: tipo de gráfico (ej. barras, líneas, circular, diagrama de flujo, imagen de un objeto), datos o rangos de valores, etiquetas de ejes, elementos clave, propósito del gráfico y cómo se relaciona con la pregunta. Si es NO, escribe N/A.]

        --- FORMATO ESPERADO DE SALIDA ---
        PREGUNTA: [Redacta aquí el enunciado de la pregunta]
        A. [Opción A]  
        B. [Opción B]  
        C. [Opción C]  
        RESPUESTA CORRECTA: [Letra de la opción correcta, por ejemplo: B]
        JUSTIFICACIONES:  
        A. [Explica por qué A es incorrecta o correcta]  
        B. [Explica por qué B es incorrecta o correcta]  
        C. [Explica por qué C es incorrecta o correcta]  
        GRAFICO_NECESARIO: [SÍ/NO]
        DESCRIPCION_GRAFICO: [Descripción detallada o N/A]
        """
        
        # Si no es el primer intento, añade las observaciones de auditoría para refinamiento
        if attempt > 1:
            prompt_content_for_llm += f"""
            --- RETROALIMENTACIÓN DE AUDITORÍA PARA REFINAMIENTO ---
            El ítem anterior no cumplió con todos los criterios. Por favor, revisa las siguientes observaciones y mejora el ítem para abordarlas.
            Observaciones del Auditor:
            {audit_observations}
            ---------------------------------------------------
            """
            # Agrega el ítem anterior para que el LLM lo pueda reformular
            prompt_content_for_llm += f"""
            --- ÍTEM ANTERIOR A REFINAR ---
            {current_item_text}
            -------------------------------
            """

        try:
            with st.spinner(f"Generando contenido con IA ({gen_model_type} - {gen_model_name}, Intento {attempt})..."):
                full_llm_response = generar_texto_con_llm(gen_model_type, gen_model_name, prompt_content_for_llm)
                
                if full_llm_response is None: # Si hubo un error en la generación con LLM
                    st.error(f"Fallo en la generación de texto con {gen_model_type} ({gen_model_name}).")
                    auditoria_status = "❌ RECHAZADO (Error de Generación)"
                    audit_observations = "El modelo de generación no pudo producir una respuesta válida."
                    break # Salir del bucle de refinamiento
                
                # --- Parsear la respuesta para extraer el ítem y la información del gráfico ---
                item_and_graphic_match = re.search(r"(PREGUNTA:.*?)(GRAFICO_NECESARIO:\s*(SÍ|NO).*?DESCRIPCION_GRAFICO:.*)", full_llm_response, re.DOTALL)
                
                if item_and_graphic_match:
                    current_item_text = item_and_graphic_match.group(1).strip()
                    grafico_info_block = item_and_graphic_match.group(2).strip()
                    
                    grafico_necesario_match = re.search(r"GRAFICO_NECESARIO:\s*(SÍ|NO)", grafico_info_block)
                    if grafico_necesario_match:
                        grafico_necesario = grafico_necesario_match.group(1).strip()

                    descripcion_grafico_match = re.search(r"DESCRIPCION_GRAFICO:\s*(.*)", grafico_info_block, re.DOTALL)
                    if descripcion_grafico_match:
                        descripcion_grafico = descripcion_grafico_match.group(1).strip()
                        if descripcion_grafico.upper() == 'N/A':
                            descripcion_grafico = ""
                else:
                    current_item_text = full_llm_response
                    grafico_necesario = "NO"
                    descripcion_grafico = ""
                    st.warning("No se pudo parsear el formato de gráfico de la respuesta. Asumiendo que no requiere gráfico.")

                st.subheader(f"Ítem Generado/Refinado (Intento {attempt}):")
                st.markdown(current_item_text)
                if grafico_necesario == "SÍ":
                    st.info(f"**Gráfico Necesario:** SÍ")
                    st.markdown(f"**Descripción del Gráfico:**\n{descripcion_grafico}")
                else:
                    st.info("**Gráfico Necesario:** NO")
                st.markdown("---")
            
            with st.spinner(f"Auditando ítem ({audit_model_type} - {audit_model_name}, Intento {attempt})..."):
                auditoria_resultado = auditar_item_con_llm(
                    audit_model_type, audit_model_name,
                    item_generado=current_item_text,
                    grado=grado_elegido, area=area_elegida, asignatura=asignatura_elegida, estacion=estacion_elegida,
                    proceso_cognitivo=proceso_cognitivo_elegido, nanohabilidad=nanohabilidad_elegida,
                    microhabilidad=microhabilidad_elegida, competencia_nanohabilidad=competencia_nanohabilidad_elegida,
                    contexto_educativo=contexto_educativo, manual_reglas_texto=manual_reglas_texto,
                    descripcion_bloom=descripcion_bloom,
                    grafico_necesario=grafico_necesario,
                    descripcion_grafico=descripcion_grafico
                )
                if auditoria_resultado is None: # Si hubo un error en la auditoría con LLM
                    st.error(f"Fallo en la auditoría con {audit_model_type} ({audit_model_name}).")
                    auditoria_status = "❌ RECHAZADO (Error de Auditoría)"
                    audit_observations = "El modelo de auditoría no pudo producir una respuesta válida."
                    break # Salir del bucle de refinamiento

                st.subheader("Resultado de Auditoría:")
                st.markdown(auditoria_resultado)
                st.markdown("---")

            # --- Extraer DICTAMEN FINAL de forma más robusta ---
            dictamen_final_match = re.search(r"DICTAMEN FINAL:\s*\[(.*?)]", auditoria_resultado, re.DOTALL)
            if dictamen_final_match:
                auditoria_status = dictamen_final_match.group(1).strip()
            else:
                auditoria_status = "❌ RECHAZADO (no se pudo extraer dictamen)"
            
            observaciones_start = auditoria_resultado.find("OBSERVACIONES FINALES:")
            if observaciones_start != -1:
                audit_observations = auditoria_resultado[observaciones_start + len("OBSERVACIONES FINALES:"):].strip()
            else:
                audit_observations = "No se pudieron extraer observaciones específicas del auditor. Posiblemente un error de formato en la respuesta del auditor."
            
            st.info(f"Dictamen extraído: {auditoria_status}. Observaciones: {audit_observations[:100]}...")

            # Guardar los datos del ítem, incluyendo el estado final de la auditoría y observaciones
            item_final_data = {
                "item_text": current_item_text,
                "classification": classification_details,
                "grafico_necesario": grafico_necesario,
                "descripcion_grafico": descripcion_grafico,
                "final_audit_status": auditoria_status, # Guarda el estado final del intento
                "final_audit_observations": audit_observations # Guarda las observaciones del intento
            }

            if auditoria_status == "✅ CUMPLE TOTALMENTE":
                st.success(f"¡El ítem ha sido auditado y CUMPLE TOTALMENTE en el intento {attempt}!")
                break # Sale del ciclo de refinamiento si es aprobado
            else:
                st.warning(f"El ítem necesita refinamiento. Dictamen: {auditoria_status}. Intentando de nuevo...")

        except Exception as e:
            st.error(f"Error durante la generación o auditoría (intento {attempt}): {e}")
            audit_observations = f"Error técnico durante la generación: {e}. Por favor, corrige este problema."
            auditoria_status = "❌ RECHAZADO (error técnico)" 
            item_final_data = {
                "item_text": current_item_text if current_item_text else "No se pudo generar el ítem debido a un error técnico.",
                "classification": classification_details,
                "grafico_necesario": "NO",
                "descripcion_grafico": "",
                "final_audit_status": auditoria_status,
                "final_audit_observations": audit_observations
            }
            break # Sale del ciclo si hay un error técnico grave

    if item_final_data is None: 
        st.error(f"No se pudo generar ningún ítem después de {max_refinement_attempts} intentos debido a fallas en la generación/auditoría.")
        return [] # Retorna una lista vacía si no se logró generar nada en absoluto.

    return [item_final_data] # Siempre devuelve una lista con el último ítem procesado.

# --- Función para exportar preguntas a un documento Word ---
def exportar_a_word(preguntas_procesadas_list):
    """
    Exporta una lista de preguntas procesadas a un documento de Word (.docx) en memoria,
    incluyendo sus detalles de clasificación, la descripción del gráfico si aplica,
    y el dictamen final de la auditoría.
    Returns: BytesIO object of the document.
    """
    doc = docx.Document()
    
    doc.add_heading('Preguntas Generadas y Auditadas', level=1)
    doc.add_paragraph('Este documento contiene los ítems generados por el sistema de IA y sus resultados de auditoría.')
    doc.add_paragraph('') # Espacio en blanco

    if not preguntas_procesadas_list:
        doc.add_paragraph('No se procesaron ítems para este informe.')

    for i, item_data in enumerate(preguntas_procesadas_list):
        pregunta_texto = item_data["item_text"]
        classification = item_data["classification"]
        grafico_necesario = item_data.get("grafico_necesario", "NO")
        descripcion_grafico = item_data.get("descripcion_grafico", "")
        final_audit_status = item_data.get("final_audit_status", "N/A")
        final_audit_observations = item_data.get("final_audit_observations", "No hay observaciones finales de auditoría.")

        doc.add_heading(f'Ítem #{i+1}', level=2)
        
        # Añadir detalles de clasificación
        doc.add_paragraph('--- Clasificación del Ítem ---') # Usando un estilo simple
        for key, value in classification.items():
            p = doc.add_paragraph()
            run = p.add_run(f"{key}: ")
            run.bold = True
            p.add_run(str(value)) # Asegurar que el valor sea string

        doc.add_paragraph('') # Espaciador
        
        # Añadir el texto de la pregunta y su formato
        lines = pregunta_texto.split('\n')
        for line in lines:
            line = line.strip() # Limpiar espacios en blanco al inicio/final
            if not line: # Saltar líneas vacías
                continue

            if line.startswith("PREGUNTA:"):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = docx.shared.Pt(12) # Opcional: fuente más grande para la pregunta
            elif line.startswith("A.") or line.startswith("B.") or line.startswith("C."):
                p = doc.add_paragraph(line)
                p.paragraph_format.left_indent = docx.shared.Inches(0.5) # Indentar opciones
            elif line.startswith("RESPUESTA CORRECTA:"):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
            elif line.startswith("JUSTIFICACIONES:"):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
            # No incluir GRAFICO_NECESARIO ni DESCRIPCION_GRAFICO directamente aquí, ya que se añaden aparte
            # y estas líneas podrían venir de la respuesta cruda de Gemini.
            elif line.startswith("GRAFICO_NECESARIO:") or line.startswith("DESCRIPCION_GRAFICO:"):
                continue # Saltar estas líneas ya que las manejamos por separado
            elif line.startswith("VALIDACIÓN DE CRITERIOS:") or line.startswith("DICTAMEN FINAL:") or line.startswith("OBSERVACIONES FINALES:"):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
            elif line.startswith("✅") or line.startswith("⚠️") or line.startswith("❌"):
                p = doc.add_paragraph(line)
                p.paragraph_format.left_indent = docx.shared.Inches(0.25) # Indentar estado de auditoría
            else:
                doc.add_paragraph(line)
        
        # Añadir descripción del gráfico si es necesario
        if grafico_necesario == "SÍ" and descripcion_grafico:
            doc.add_paragraph('')
            p = doc.add_paragraph()
            run = p.add_run("--- Gráfico Sugerido ---")
            run.bold = True
            doc.add_paragraph(f"**Tipo y Descripción del Gráfico:** {descripcion_grafico}")
            doc.add_paragraph('') # Espacio adicional

        # Añadir el dictamen final y las observaciones de la auditoría para CADA ítem
        doc.add_paragraph('')
        p = doc.add_paragraph()
        run = p.add_run("--- Resultado Final de Auditoría ---")
        run.bold = True
        doc.add_paragraph(f"**DICTAMEN FINAL:** {final_audit_status}")
        doc.add_paragraph(f"**OBSERVACIONES FINALES:** {final_audit_observations}")
        doc.add_paragraph('') # Espacio adicional

        doc.add_page_break() # Separar cada pregunta con un salto de página

    # Guardar el documento en un buffer en memoria
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0) # Regresar al inicio del buffer
    return buffer

# --- Interfaz de Usuario de Streamlit ---
st.title("📚 Generador y Auditor de Ítems Educativos con IA 🧠")
st.markdown("Esta aplicación genera ítems de selección múltiple basados en tus especificaciones y los audita automáticamente.")

# --- Sección de Carga de Archivos ---
st.sidebar.header("Carga de Archivos")
uploaded_excel_file = st.sidebar.file_uploader("Sube tu archivo Excel (ESTRUCTURA_TOTAL.xlsx)", type=["xlsx"])
uploaded_pdf_file = st.sidebar.file_uploader("Sube tu archivo PDF (Manual_construccion_pruebas_IMProve.pdf)", type=["pdf"])

df_datos = None
manual_reglas_texto = ""

if uploaded_excel_file:
    df_datos = leer_excel_cargado(uploaded_excel_file)

if uploaded_pdf_file:
    manual_reglas_texto = leer_pdf_cargado(uploaded_pdf_file)
    max_manual_length = 15000 
    if len(manual_reglas_texto) > max_manual_length:
        st.sidebar.warning(f"Manual es demasiado largo ({len(manual_reglas_texto)} caracteres). Truncando a {max_manual_length} caracteres para la IA.")
        manual_reglas_texto = manual_reglas_texto[:max_manual_length]
    st.sidebar.info(f"Manual de reglas cargado. Longitud final: {len(manual_reglas_texto)} caracteres.")

# --- Selección de Modelos ---
st.sidebar.header("Configuración de Modelos de IA")

# Generador
st.sidebar.subheader("Modelo para Generación de Ítems")
gen_model_type = st.sidebar.radio("Tipo de Modelo", ["Gemini", "GPT"], key="gen_model_type")
gen_model_name = ""
if gen_model_type == "Gemini":
    gen_model_name = st.sidebar.selectbox("Nombre del Modelo Gemini", ["gemini-1.5-flash", "gemini-1.5-pro"], key="gen_gemini_name")
else: # GPT
    gen_model_name = st.sidebar.selectbox("Nombre del Modelo GPT", ["gpt-4o", "gpt-4-turbo", "gpt-3.5-turbo"], key="gen_gpt_name")

# Auditor
st.sidebar.subheader("Modelo para Auditoría de Ítems")
audit_model_type = st.sidebar.radio("Tipo de Modelo", ["Gemini", "GPT"], key="audit_model_type")
audit_model_name = ""
if audit_model_type == "Gemini":
    audit_model_name = st.sidebar.selectbox("Nombre del Modelo Gemini", ["gemini-1.5-flash", "gemini-1.5-pro"], key="audit_gemini_name")
else: # GPT
    audit_model_name = st.sidebar.selectbox("Nombre del Modelo GPT", ["gpt-4o", "gpt-4-turbo", "gpt-3.5-turbo"], key="audit_gpt_name")


# --- Lógica Principal de la Aplicación ---
if df_datos is not None and (gemini_config_ok or openai_config_ok):
    st.header("Selecciona los Criterios para la Generación")

    # Obtener valores únicos para cada columna para los selectbox
    all_grades = df_datos['GRADO'].dropna().unique().tolist()
    grado_seleccionado = st.selectbox("Grado", sorted(all_grades), key="grado_sel")

    # Filtrar el DataFrame según la selección del grado
    df_filtrado_grado = df_datos[df_datos['GRADO'].astype(str).str.upper() == str(grado_seleccionado).upper()]
    all_areas = df_filtrado_grado['ÁREA'].dropna().unique().tolist()
    area_seleccionada = st.selectbox("Área", sorted(all_areas), key="area_sel")

    # Filtrar según la selección del área
    df_filtrado_area = df_filtrado_grado[df_filtrado_grado['ÁREA'].astype(str).str.upper() == str(area_seleccionada).upper()]
    all_asignaturas = df_filtrado_area['ASIGNATURA'].dropna().unique().tolist()
    asignatura_seleccionada = st.selectbox("Asignatura", sorted(all_asignaturas), key="asignatura_sel")

    # Filtrar según la selección de asignatura
    df_filtrado_asignatura = df_filtrado_area[df_filtrado_area['ASIGNATURA'].astype(str).str.upper() == str(asignatura_seleccionada).upper()]
    all_estaciones = df_filtrado_asignatura['ESTACIÓN'].dropna().unique().tolist()
    estacion_seleccionada = st.selectbox("Estación", sorted(all_estaciones), key="estacion_sel")

    # Filtrar según la selección de estación
    df_filtrado_estacion = df_filtrado_asignatura[df_filtrado_asignatura['ESTACIÓN'].astype(str).str.upper() == str(estacion_seleccionada).upper()]
    all_procesos = df_filtrado_estacion['PROCESO COGNITIVO'].dropna().unique().tolist()
    proceso_cognitivo_seleccionado = st.selectbox("Proceso Cognitivo", sorted(all_procesos), key="proceso_sel")

    # Filtrar según la selección de proceso cognitivo
    df_filtrado_proceso = df_filtrado_estacion[df_filtrado_estacion['PROCESO COGNITIVO'].astype(str).str.upper() == str(proceso_cognitivo_seleccionado).upper()]
    all_nanohabilidades = df_filtrado_proceso['NANOHABILIDAD'].dropna().unique().tolist()
    nanohabilidad_seleccionada = st.selectbox("Nanohabilidad", sorted(all_nanohabilidades), key="nanohabilidad_sel")

    # Después de todas las selecciones, se filtra el DataFrame final
    df_item_seleccionado = df_filtrado_proceso[df_filtrado_proceso['NANOHABILIDAD'].astype(str).str.upper() == str(nanohabilidad_seleccionada).upper()]

    if df_item_seleccionado.empty:
        st.warning("No se encontraron datos en el Excel para la combinación de criterios seleccionada.")
        st.stop() # Detener la ejecución aquí si no hay datos

    # --- Información Adicional del Usuario ---
    st.subheader("Información Adicional para el Ítem")
    opcion_info_adicional = st.radio(
        "¿Deseas proporcionar alguna información adicional o contexto para la generación del ítem?",
        ("No", "Sí"),
        key="info_ad_radio"
    )
    informacion_adicional_usuario = ""
    if opcion_info_adicional == "Sí":
        informacion_adicional_usuario = st.text_area("Escribe la información adicional que deseas incluir:", key="info_ad_text")

    # --- Botón para Generar y Auditar ---
    if st.button("Generar y Auditar Ítem"):
        if df_item_seleccionado.empty:
            st.error("Por favor, selecciona criterios válidos que resulten en datos para generar el ítem.")
        elif (gen_model_type == "Gemini" and not gemini_config_ok) or (gen_model_type == "GPT" and not openai_config_ok):
            st.error(f"Por favor, configura la API Key para el modelo de generación ({gen_model_type}).")
        elif (audit_model_type == "Gemini" and not gemini_config_ok) or (audit_model_type == "GPT" and not openai_config_ok):
            st.error(f"Por favor, configura la API Key para el modelo de auditoría ({audit_model_type}).")
        else:
            st.markdown("---")
            st.info("Iniciando generación y auditoría del ítem. Esto puede tardar unos momentos...")

            # Preparar los criterios de generación para la función
            criterios_para_preguntas = {
                "tipo_pregunta": "opción múltiple con 3 opciones", 
                "dificultad": "media", # Se podría hacer seleccionable también
                "num_preguntas": 1, 
                "contexto_educativo": "estudiantes de preparatoria (bachillerato)", # Se podría hacer seleccionable
                "formato_justificacion": """
                    • Justificación correcta: debe explicar el razonamiento o proceso cognitivo (NO por descarte).
                    • Justificaciones incorrectas: deben redactarse como: “El estudiante podría escoger la opción X porque… Sin embargo, esto es incorrecto porque…”
                """
            }

            # Llamar a la función para generar y auditar el ítem
            item_procesado_individual = generar_pregunta_con_seleccion( # Se actualiza el nombre de la función
                gen_model_type, gen_model_name, audit_model_type, audit_model_name, # Pasa los tipos y nombres de modelos
                fila_datos=df_item_seleccionado.iloc[0], 
                criterios_generacion=criterios_para_preguntas,
                manual_reglas_texto=manual_reglas_texto,
                informacion_adicional_usuario=informacion_adicional_usuario
            )

            # Almacenar el resultado del procesamiento en el estado de la sesión
            if item_procesado_individual: # Si se procesó y obtuvo un resultado (aprobado o no)
                st.session_state['last_processed_item_data'] = item_procesado_individual[0] # Guardamos el diccionario directamente
                
                if item_procesado_individual[0].get('final_audit_status') == "✅ CUMPLE TOTALMENTE":
                    st.success("¡Ítem generado y aprobado por el auditor! Listo para exportar.")
                else:
                    st.warning(f"Ítem generado pero NO aprobado por el auditor. Dictamen final: {item_procesado_individual[0].get('final_audit_status')}. Se guardará la última versión con observaciones.")
                
                st.subheader("Último Ítem Procesado:")
                st.markdown(item_procesado_individual[0]['item_text'])
                st.write("--- Clasificación ---")
                for key, value in item_procesado_individual[0]['classification'].items():
                    st.write(f"- **{key}**: {value}")
                
                if item_procesado_individual[0]['grafico_necesario'] == "SÍ":
                    st.write("--- Gráfico Sugerido ---")
                    st.write(f"**Descripción del Gráfico:** {item_procesado_individual[0]['descripcion_grafico']}")
                
                st.write("--- Resultado Final de Auditoría ---")
                st.write(f"**DICTAMEN FINAL:** {item_procesado_individual[0]['final_audit_status']}")
                st.write(f"**OBSERVACIONES FINALES:** {item_procesado_individual[0]['final_audit_observations']}")
                st.markdown("---")

            else: # Si la función generador_preguntas_con_llm devolvió una lista vacía o None
                st.error("No se pudo generar ni procesar el ítem. Verifica tus entradas y la conexión a la IA.")
                st.session_state['last_processed_item_data'] = None # No hay ítem para guardar
    
    # --- Sección de Exportación a Word (Siempre visible al final) ---
    st.header("Exportar a Documento Word")

    if 'last_processed_item_data' in st.session_state and st.session_state['last_processed_item_data'] is not None:
        st.write("Hay un ítem procesado disponible para exportar (aprobado o la última versión con observaciones).")
        nombre_archivo_word = st.text_input("Ingresa el nombre deseado para el archivo Word (sin la extensión .docx):", key="word_filename")
        
        if nombre_archivo_word:
            # Creamos una lista con el único ítem procesado (aprobado o no) para la función exportar_a_word
            items_para_exportar = [st.session_state['last_processed_item_data']]
            
            word_buffer = exportar_a_word(items_para_exportar)
            
            st.download_button(
                label="Descargar Documento Word",
                data=word_buffer,
                file_name=f"{nombre_archivo_word}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingprocessingml.document"
            )
            st.info("Haz clic en el botón de arriba para descargar tu archivo Word. Se guardará en la carpeta de descargas de tu navegador.")
        else:
            st.warning("Por favor, ingresa un nombre para el archivo Word para habilitar la descarga.")
    else:
        st.info("No hay ítems procesados disponibles para exportar a Word en este momento.")
        st.write("Genera y audita un ítem para que esté disponible aquí.")

elif uploaded_excel_file is None:
    st.info("Por favor, sube tu archivo Excel para comenzar.")
elif not (gemini_config_ok or openai_config_ok):
    st.info("Por favor, ingresa al menos una API Key de Gemini o OpenAI en la barra lateral para comenzar.")
